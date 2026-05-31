from docx import Document
from docx.shared import Pt
import re

md_path = 'docs/ARYAN_PRICE_ACTION_INDICATOR_PAPER.md'
docx_path = 'docs/ARYAN_PRICE_ACTION_INDICATOR_PAPER.docx'

def is_table_line(line):
    return line.strip().startswith('|') and line.strip().endswith('|')

def parse_table(lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def main():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        if not line.strip():
            i += 1
            continue
        # headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            lvl = len(m.group(1))
            text = m.group(2).strip()
            # Map to docx heading levels (1..4)
            heading_lvl = min(lvl, 4)
            doc.add_heading(text, level=heading_lvl)
            i += 1
            continue
        # horizontal rule
        if re.match(r'^---+$', line.strip()):
            doc.add_paragraph()
            i += 1
            continue
        # table
        if is_table_line(line):
            tbl_lines = [line]
            j = i+1
            # include separator and subsequent table rows
            while j < len(lines) and is_table_line(lines[j]):
                tbl_lines.append(lines[j].rstrip('\n'))
                j += 1
            # remove possible separator like |---|
            tbl_clean = [l for l in tbl_lines if not re.match(r'^\|?\s*-{3,}\s*\|', l)]
            if tbl_clean:
                rows = parse_table(tbl_clean)
                # create table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        table.rows[r_idx].cells[c_idx].text = cell
            i = j
            continue
        # figure placeholder
        if '[INSERT FIGURE' in line or line.strip().startswith('Figure') and 'INSERT' in line:
            p = doc.add_paragraph()
            p.alignment = 1  # center
            run = p.add_run(line.strip())
            run.italic = True
            i += 1
            continue
        # bullet lists
        if re.match(r'^\s*[-\*]\s+.*', line):
            # collect contiguous bullets
            bullets = []
            j = i
            while j < len(lines) and re.match(r'^\s*[-\*]\s+.*', lines[j]):
                bullets.append(re.sub(r'^\s*[-\*]\s+', '', lines[j].strip()))
                j += 1
            for b in bullets:
                doc.add_paragraph(b, style='List Bullet')
            i = j
            continue
        # numbered lists
        if re.match(r'^\s*\d+\.\s+.*', line):
            nums = []
            j = i
            while j < len(lines) and re.match(r'^\s*\d+\.\s+.*', lines[j]):
                nums.append(re.sub(r'^\s*\d+\.\s+', '', lines[j].strip()))
                j += 1
            for n in nums:
                doc.add_paragraph(n, style='List Number')
            i = j
            continue
        # normal paragraph
        # accumulate consecutive non-empty lines until blank or special
        para_lines = [line]
        j = i+1
        while j < len(lines) and lines[j].strip() and not re.match(r'^(#{1,6})\s+', lines[j]) and not is_table_line(lines[j]) and not re.match(r'^\s*[-\*]\s+', lines[j]) and not re.match(r'^\s*\d+\.\s+', lines[j]):
            para_lines.append(lines[j].rstrip('\n'))
            j += 1
        paragraph_text = ' '.join([l.strip() for l in para_lines])
        doc.add_paragraph(paragraph_text)
        i = j

    doc.save(docx_path)
    print('Saved', docx_path)

if __name__ == '__main__':
    main()
